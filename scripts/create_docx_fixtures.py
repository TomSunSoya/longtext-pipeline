#!/usr/bin/env python3
"""Generate test DOCX fixtures for longtext-pipeline."""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_UNDERLINE


def create_simple_docx(filepath):
    """Create a simple DOCX without special features."""
    doc = Document()

    doc.add_heading("Simple Document", 0)

    doc.add_paragraph(
        "This is a simple document for testing basic DOCX parsing functionality. "
        "It contains plain text without complex formatting, tables, or images."
    )

    for i in range(1, 6):
        doc.add_heading(f"Section {i}: Basic Content", level=1)
        doc.add_paragraph(
            f"This is paragraph {i} in the simple document. "
            "It demonstrates basic text formatting and structure."
        )
        doc.add_paragraph(
            "Another paragraph with more text to ensure the document is reasonable in size. "
            "This paragraph tests paragraph boundaries and text extraction."
        )

    doc.save(filepath)


def create_with_tables_docx(filepath):
    """Create a DOCX with tables."""
    doc = Document()

    doc.add_heading("Document with Tables", 0)

    # Simple table
    doc.add_heading("Simple Table", level=1)
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Column A"
    hdr_cells[1].text = "Column B"
    hdr_cells[2].text = "Column C"

    # Data rows
    for i in range(1, 4):
        row_cells = table.rows[i].cells
        row_cells[0].text = f"Row {i} Col 1"
        row_cells[1].text = f"Row {i} Col 2"
        row_cells[2].text = f"Row {i} Col 3"

    doc.add_paragraph(
        "The table above contains structured data that tests table parsing functionality."
    )

    # Another nested table
    doc.add_heading("Nested Table Structure", level=1)
    table2 = doc.add_table(rows=3, cols=2)
    table2.style = "Table Grid"

    table2.rows[0].cells[0].text = "Main Category"
    table2.rows[0].cells[1].text = "Details"

    inner_table = table2.rows[1].cells[0].add_table(rows=3, cols=2)
    inner_table.rows[0].cells[0].text = "Sub"
    inner_table.rows[0].cells[1].text = "Item"
    inner_table.rows[1].cells[0].text = "Item A"
    inner_table.rows[1].cells[1].text = "Data 1"
    inner_table.rows[2].cells[0].text = "Item B"
    inner_table.rows[2].cells[1].text = "Data 2"

    table2.rows[1].cells[1].text = "Complex nested structure here"
    table2.rows[2].cells[0].text = "Bottom"
    table2.rows[2].cells[1].text = "End content"

    doc.save(filepath)


def create_tracked_changes_docx(filepath):
    """Create a DOCX that simulates tracked changes (as a test fixture)."""
    doc = Document()

    doc.add_heading("Document with Tracked Changes", 0)

    # Mark this document as having tracked changes by using special markers
    # The actual tracked changes XML is complex, so we simulate with markers
    doc.add_paragraph(
        "NOTE: This document contains content that simulates tracked changes. "
        "In real documents processed by longtext-pipeline, tracked changes appear as "
        "insertions (text marked with <w:ins>) and deletions (text marked with <w:del>)."
    )

    doc.add_heading("Original Content", level=1)
    doc.add_paragraph(
        "This text represents what was in the original document before edits."
    )

    doc.add_heading("Insertions (simulated)", level=1)
    doc.add_paragraph("[INSERTED] This text was added via tracked changes feature.")
    doc.add_paragraph("Normal text continues here after the insertion marker.")

    doc.add_heading("Deletions (simulated)", level=1)
    doc.add_paragraph("[DELETED] This content was removed via tracked changes.")
    doc.add_paragraph("This text would appear where the deletion was in the original.")

    doc.add_paragraph(
        "For testing purposes, track changes are often stripped or marked "
        "separately during document processing."
    )

    doc.save(filepath)


def create_with_images_docx(filepath):
    """Create a DOCX with embedded images."""
    doc = Document()

    doc.add_heading("Document with Images", 0)

    # Add a title image placeholder (using a simple colored shape)
    paragraph = doc.add_paragraph()
    paragraph.alignment = 1  # Center
    run = paragraph.add_run("Image Placeholder 1")
    run.font.color.rgb = RGBColor(100, 100, 100)

    # Add some spacer text
    doc.add_paragraph(
        "This document contains image placeholders for testing image handling."
    )

    # Add figure caption
    doc.add_paragraph("Figure 1: Example visualization", style="Caption")

    doc.add_paragraph(
        "The images in this document could be charts, screenshots, or diagrams "
        "that require separate processing from the text content."
    )

    # Another image placeholder
    paragraph2 = doc.add_paragraph()
    paragraph2.alignment = 1
    run2 = paragraph2.add_run("Image Placeholder 2 - Chart Visualization")
    run2.font.color.rgb = RGBColor(80, 80, 80)

    doc.save(filepath)


def create_complex_formatting_docx(filepath):
    """Create a DOCX with complex formatting (styles, colors, multiple levels)."""
    doc = Document()

    doc.add_heading("Complex Formatting Document", 0)

    # Different heading levels
    doc.add_heading("Level 1 Heading", level=1)
    doc.add_paragraph("This paragraph uses the default style with standard formatting.")

    doc.add_heading("Level 2 Heading", level=2)
    doc.add_paragraph(
        "This is a paragraph under a level 2 heading with different visual hierarchy."
    )

    doc.add_heading("Level 3 Heading", level=3)
    doc.add_paragraph("Content under level 3 appears indented and styled differently.")

    # Bold and italic text
    doc.add_heading("Styling Examples", level=2)
    para = doc.add_paragraph()
    run1 = para.add_run("This is bold text. ")
    run1.bold = True
    run2 = para.add_run("This is italic text. ")
    run2.italic = True
    run3 = para.add_run("This is underlined text. ")
    run3.underline = WD_UNDERLINE.SINGLE
    para.add_run("And this is normal text again.")

    # Font color and size variations
    doc.add_heading("Color and Size Variations", level=2)
    para2 = doc.add_paragraph()
    run4 = para2.add_run("Red text example. ")
    run4.font.color.rgb = RGBColor(255, 0, 0)
    run4.font.size = Pt(14)

    run5 = para2.add_run("Blue text example. ")
    run5.font.color.rgb = RGBColor(0, 0, 255)
    run5.font.size = Pt(12)

    run6 = para2.add_run("Smaller green text. ")
    run6.font.color.rgb = RGBColor(0, 128, 0)
    run6.font.size = Pt(10)

    # Numbered list
    doc.add_heading("Lists and Outlines", level=2)
    doc.add_paragraph("First item in numbered list", style="List Number")
    doc.add_paragraph("Second item in numbered list", style="List Number")
    doc.add_paragraph("Third item in numbered list", style="List Number")

    # Bulleted list
    doc.add_paragraph("Bullet item one", style="List Bullet")
    doc.add_paragraph("Bullet item two", style="List Bullet")
    doc.add_paragraph("Bullet item three", style="List Bullet")

    # Block quote
    doc.add_paragraph(
        "This is a block quote example for testing quote detection.",
        style="Intense Quote",
    )

    doc.save(filepath)


def main():
    """Generate all DOCX fixtures."""
    output_dir = Path("tests/fixtures/docx")
    output_dir.mkdir(parents=True, exist_ok=True)

    fixtures = [
        ("01_simple.docx", create_simple_docx),
        ("02_with_tables.docx", create_with_tables_docx),
        ("03_tracked_changes.docx", create_tracked_changes_docx),
        ("04_with_images.docx", create_with_images_docx),
        ("05_complex_formatting.docx", create_complex_formatting_docx),
    ]

    for filename, creator in fixtures:
        filepath = output_dir / filename
        creator(str(filepath))
        size = filepath.stat().st_size / 1024  # KB
        print(f"  {filename}: {size:.1f} KB")

    print(f"\nGenerated {len(fixtures)} DOCX fixtures")


if __name__ == "__main__":
    main()
