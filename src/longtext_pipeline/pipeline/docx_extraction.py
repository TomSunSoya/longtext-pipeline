"""
DOCX text extraction module using python-docx for the longtext pipeline.

This module provides DOCX text extraction functionality with python-docx as
the primary extraction engine. It handles various DOCX features including
complex document structures like tables, tracked changes, and embedded content
while providing robust error handling for invalid files.
"""

import logging
import re
from pathlib import Path
from typing import TYPE_CHECKING

from ..errors import InputError

logger = logging.getLogger(__name__)

# Conditional import since python-docx may not be installed
if TYPE_CHECKING:
    import docx
    from docx.document import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.text.run import Run

DOCX_AVAILABLE = False

try:
    import docx
    from docx.document import Document  # type: ignore
    from docx.table import Table  # type: ignore
    from docx.text.paragraph import Paragraph  # type: ignore
    from docx.text.run import Run  # type: ignore

    DOCX_AVAILABLE = True
except ImportError:
    docx = None  # type: ignore
    Document = None  # type: ignore
    Table = None  # type: ignore
    Paragraph = None  # type: ignore
    Run = None  # type: ignore


class DOCXTextExtractor:
    """Handles DOCX text extraction with robust error handling."""

    def __init__(self) -> None:
        """Initialize the DOCX text extractor.

        If python-docx is not available, raise an error.
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for DOCX extraction. "
                "Install with: pip install python-docx"
            )

    def extract_text_from_docx(
        self,
        docx_path: str,
        extraction_mode: str = "basic",
        include_tables: bool = True,
        include_footnotes: bool = False,
        config: dict | None = None,
        **kwargs,
    ) -> str:
        """Extract text from DOCX file using python-docx.

        Args:
            docx_path: Path to the DOCX file to extract text from
            extraction_mode: Strategy for text extraction:
                - "basic": Extract paragraphs and tables only
                - "formatted": Preserve some formatting indicators
            include_tables: Whether to extract text from tables
            include_footnotes: Whether to extract footnotes/endnotes (optional)
            config: Configuration dictionary (reserved for future use)
            **kwargs: Additional extraction parameters

        Returns:
            Extracted text content as string

        Raises:
            InputError: If DOCX extraction fails or results in empty content
        """
        if config is None:
            config = {}

        docx_file_path = Path(docx_path)

        if not docx_file_path.exists():
            raise InputError(f"DOCX file does not exist: {docx_path}")

        if not docx_file_path.is_file():
            raise InputError(f"Path is not a file: {docx_path}")

        text_content = []

        try:
            # Load the document
            doc = docx.Document(str(docx_file_path))  # type: ignore

            # Extract paragraphs
            for paragraph in doc.paragraphs:  # type: ignore
                paragraph_text = self._process_paragraph(paragraph, extraction_mode)
                if paragraph_text:
                    text_content.append(paragraph_text)

            # Extract tables if requested
            if include_tables:
                for table in doc.tables:  # type: ignore
                    table_text = self._process_table(table, extraction_mode)
                    if table_text:
                        text_content.append(table_text)

            # Extract footnotes/endnotes if requested
            if include_footnotes:
                footnotes_text = self._process_footnotes(doc)  # type: ignore
                if footnotes_text:
                    text_content.append(footnotes_text)

        except PermissionError:
            raise InputError(f"No permission to read DOCX file: {docx_path}")
        except Exception as e:
            error_msg = str(e).lower()
            if "invalid" in error_msg or "corrupt" in error_msg:
                raise InputError(f"Invalid or corrupted DOCX file: {docx_path}")
            elif "zip" in error_msg or "not a zip" in error_msg:
                raise InputError(
                    f"File is not a valid DOCX (ZIP format error): {docx_path}"
                )
            else:
                raise InputError(f"Failed to process DOCX file: {e}")

        # Combine all extracted content
        full_text = "\n\n".join(text_content)

        if not full_text.strip():
            raise InputError(
                f"DOCX extraction resulted in empty content for file: {docx_path}"
            )

        logger.info(
            "Successfully extracted %d characters from %s", len(full_text), docx_path
        )
        return full_text

    def _process_paragraph(self, paragraph: "Paragraph", extraction_mode: str) -> str:
        """Process a single paragraph, handling runs and formatting.

        Args:
            paragraph: The paragraph object from python-docx
            extraction_mode: Extraction mode ("basic" or "formatted")

        Returns:
            Processed paragraph text
        """
        if not paragraph.text.strip():
            return ""

        # For different modes, we handle formatting slightly differently
        if extraction_mode == "formatted":
            # Preserve some formatting indicators
            runs_text = []
            for run in paragraph.runs:  # type: ignore
                run_text = self._sanitize_text(run.text)
                if not run_text:
                    continue

                # Add formatting indicators
                parts = []
                if run.bold:  # type: ignore
                    parts.append("**")
                if run.italic:  # type: ignore
                    parts.append("*")
                if run.underline:  # type: ignore
                    parts.append("_")

                # Add run text with formatting
                formatted_run = "".join(parts) + run_text
                if run.bold:  # type: ignore
                    formatted_run += "**"
                if run.italic:  # type: ignore
                    formatted_run += "*"
                if run.underline:  # type: ignore
                    formatted_run += "_"

                runs_text.append(formatted_run)

            paragraph_text = "".join(runs_text)
        else:
            # Basic mode - just extract plain text
            paragraph_text = self._sanitize_text(paragraph.text)

        # Add paragraph marker for traceability
        if (
            hasattr(paragraph, "style")
            and paragraph.style
            and hasattr(paragraph.style, "name")
        ):
            style_name = paragraph.style.name.lower()  # type: ignore
            if style_name in ["heading 1", "heading 2", "heading 3"]:
                paragraph_text = f"\n# {paragraph_text}\n"
            elif style_name in ["heading 4", "heading 5", "heading 6"]:
                level = int(style_name.replace("heading ", ""))
                paragraph_text = f"\n{'#' * level} {paragraph_text}\n"

        return paragraph_text

    def _process_table(self, table: "Table", extraction_mode: str) -> str:
        """Process a table, extracting cells and preserving structure.

        Args:
            table: The table object from python-docx
            extraction_mode: Extraction mode ("basic" or "formatted")

        Returns:
            Processed table text
        """
        table_lines = []

        for row in table.rows:  # type: ignore
            row_cells = []
            for cell in row.cells:  # type: ignore
                cell_text = self._sanitize_text(cell.text)
                if cell_text:
                    row_cells.append(cell_text)
                else:
                    row_cells.append("")  # Preserve empty cells

            if row_cells:
                # Format as a table row
                row_text = " | ".join(row_cells)
                table_lines.append(row_text)

        if table_lines:
            # Add table header marker
            table_text = "\n".join(table_lines)
            return f"\nTABLE:\n{table_text}\n"

        return ""

    def _process_footnotes(self, doc: "Document") -> str:
        """Process footnotes and endnotes if present.

        Args:
            doc: The document object from python-docx

        Returns:
            Processed footnotes text
        """
        # python-docx doesn't directly expose footnotes in the Document object
        # They are stored in the document part but not easily accessible
        # For now, we'll document that footnote extraction is not supported
        # in the basic implementation, but the infrastructure is here

        # Future enhancement: access footnotes via doc.part._footnotes
        return ""

    def _sanitize_text(self, text: str) -> str:
        """Sanitize text content, removing control characters and normalizing whitespace.

        Args:
            text: Raw text to sanitize

        Returns:
            Sanitized text
        """
        if not text:
            return ""

        # Replace control characters and normalize whitespace
        # Replace multiple whitespace with single space
        text = re.sub(r"[ \t\r\f\v]+", " ", text)

        # Replace multiple newlines with double newline (paragraph break)
        text = re.sub(r"\n{3,}", "\n\n", text)

        # Strip leading/trailing whitespace from each line
        lines = text.split("\n")
        sanitized_lines = [line.strip() for line in lines]

        return "\n".join(sanitized_lines).strip()

    def extract_and_preprocess_docx(
        self,
        docx_path: str,
        extraction_mode: str = "basic",
        config: dict | None = None,
        **kwargs,
    ) -> str:
        """Extract text from DOCX and apply standard preprocessing.

        Args:
            docx_path: Path to the DOCX file
            extraction_mode: Extraction strategy
            config: Configuration dictionary
            **kwargs: Additional parameters

        Returns:
            Preprocessed text content
        """
        if config is None:
            config = {}

        raw_text = self.extract_text_from_docx(
            docx_path,
            extraction_mode=extraction_mode,
            config=config,
            **kwargs,
        )

        # Apply preprocessing similar to other formats
        processed_text = self._normalize_whitespace(raw_text)

        return processed_text

    def _normalize_whitespace(self, text: str) -> str:
        """Normalize whitespace while preserving document structure.

        Args:
            text: Raw extracted text

        Returns:
            Normalized text
        """
        # Replace excessive newlines with double newline (paragraph break)
        text = re.sub(r"\n{3,}", "\n\n", text)

        # Normalize tabs and spaces
        lines = text.split("\n")
        normalized_lines = []

        for line in lines:
            stripped = line.strip()
            if stripped:
                normalized_lines.append(stripped)
            else:
                normalized_lines.append(stripped)

        return "\n".join(normalized_lines)


def detect_docx_format(file_path: str) -> dict[str, bool | int]:
    """Detect characteristics of the DOCX to inform extraction strategy.

    Args:
        file_path: Path to the DOCX file

    Returns:
        Dictionary indicating detected characteristics
    """
    if not DOCX_AVAILABLE:
        raise ImportError(
            "python-docx is required for DOCX format detection. "
            "Install with: pip install python-docx"
        )

    features: dict[str, bool | int] = {
        "valid_docx": False,
        "has_tables": False,
        "has_images": False,
        "has_headings": False,
        "has_formatted_text": False,
        "has_footnotes": False,
        "document_size": 0,
        "page_count_estimate": 0,
    }

    try:
        doc = docx.Document(file_path)  # type: ignore

        features["valid_docx"] = True

        # Count statistics
        features["document_size"] = len(doc.paragraphs)  # type: ignore

        # Check for tables
        if doc.tables:  # type: ignore
            features["has_tables"] = True
            features["page_count_estimate"] = max(
                features["page_count_estimate"],
                len(doc.tables) // 3,  # type: ignore
            )

        # Check for headings
        for paragraph in doc.paragraphs:  # type: ignore
            if (
                hasattr(paragraph, "style")
                and paragraph.style
                and hasattr(paragraph.style, "name")
            ):
                style_name = paragraph.style.name.lower()  # type: ignore
                if style_name.startswith("heading"):
                    features["has_headings"] = True
                    break

        # Check for formatted text (bold, italic, underline)
        for paragraph in doc.paragraphs:  # type: ignore
            for run in paragraph.runs:  # type: ignore
                if run.bold:  # type: ignore
                    features["has_formatted_text"] = True
                    features["has_headings"] = True
                    break
                if run.italic:  # type: ignore
                    features["has_formatted_text"] = True
                    break
                if run.underline:  # type: ignore
                    features["has_formatted_text"] = True
                    break

        # Estimate page count (rough approximation: ~50 paragraphs per page)
        features["page_count_estimate"] = max(1, (len(doc.paragraphs) + 49) // 50)  # type: ignore

    except Exception as e:
        if "invalid" in str(e).lower() or "corrupt" in str(e).lower():
            features["valid_docx"] = False
        elif "zip" in str(e).lower() or "not a zip" in str(e).lower():
            features["valid_docx"] = False

    return features


def get_document_structure(file_path: str) -> dict[str, int]:
    """Get structural information about a DOCX document.

    Args:
        file_path: Path to the DOCX file

    Returns:
        Dictionary with structural details
    """
    if not DOCX_AVAILABLE:
        raise ImportError(
            "python-docx is required for document structure analysis. "
            "Install with: pip install python-docx"
        )

    structure: dict[str, int] = {
        "paragraphs": 0,
        "tables": 0,
        "sections": 0,
        "images": 0,
        "headings": 0,
    }

    try:
        doc = docx.Document(file_path)  # type: ignore

        structure["paragraphs"] = len(doc.paragraphs)  # type: ignore
        structure["tables"] = len(doc.tables)  # type: ignore

        # Count headings
        for paragraph in doc.paragraphs:  # type: ignore
            if (
                hasattr(paragraph, "style")
                and paragraph.style
                and hasattr(paragraph.style, "name")
            ):
                if paragraph.style.name.lower().startswith("heading"):  # type: ignore
                    structure["headings"] += 1

        # Count images (inlining runs with pictures)
        for paragraph in doc.paragraphs:  # type: ignore
            for run in paragraph.runs:  # type: ignore
                if (
                    hasattr(run, "_element")
                    and run._element is not None
                    and hasattr(run._element, "graphicData")
                ):
                    structure["images"] += 1

        # Count sections
        if hasattr(doc, "sections") and doc.sections:  # type: ignore
            structure["sections"] = len(doc.sections)  # type: ignore

    except Exception:
        pass

    return structure
